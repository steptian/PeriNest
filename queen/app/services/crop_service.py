"""Crop 嗦囊业务（RAG 知识库 ingest + 检索）。

权限逻辑只写在这里（共生体原则：REST 与 MCP 复用同一份）。
分块策略 v1：按段落聚合，~600 字一块，中文优先。
"""
import structlog
from sqlalchemy import delete, func, select
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.crop import CropChunk, CropDocument
from app.schemas.request import CropDocumentCreate
from app.services import crop_vector_store
from app.services.embedding_service import embed_texts, pack_vector

logger = structlog.get_logger(__name__)

CHUNK_TARGET_CHARS = 600  # 聚合目标块大小

# 支持的吞入文件格式（模板务实版：文字层提取，扫描件/OCR 不做——
# ack-agent 有生产级 OCR 方案，需要时再抄）
UPLOAD_MAX_BYTES = 10 * 1024 * 1024
UPLOAD_TYPES = {
    ".txt": "text",
    ".md": "markdown",
    ".markdown": "markdown",
    ".pdf": "pdf",
    ".docx": "docx",
}


class UploadUnsupported(Exception):
    """不支持的文件格式/损坏文件。"""


def extract_text(filename: str, raw: bytes) -> tuple[str, str]:
    """按扩展名提取文本。返回 (文本, source_type)。失败抛 UploadUnsupported。"""
    import io

    name = filename.lower()
    ext = next((e for e in UPLOAD_TYPES if name.endswith(e)), None)
    if ext is None:
        raise UploadUnsupported(f"不支持的格式「{filename}」——支持 {'/'.join(UPLOAD_TYPES)}")
    if len(raw) > UPLOAD_MAX_BYTES:
        raise UploadUnsupported("文件超过 10MB 限制")
    source_type = UPLOAD_TYPES[ext]

    if ext in (".txt", ".md", ".markdown"):
        try:
            return raw.decode("utf-8"), source_type
        except UnicodeDecodeError:
            try:
                return raw.decode("gbk"), source_type
            except UnicodeDecodeError as e:
                raise UploadUnsupported("文本编码无法识别（支持 UTF-8/GBK）") from e

    if ext == ".pdf":
        try:
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(raw))
            pages = [(page.extract_text() or "").strip() for page in reader.pages]
            text = "\n\n".join(p for p in pages if p)
        except Exception as e:
            raise UploadUnsupported(f"PDF 解析失败（加密/损坏）：{e}") from e
        if len(text) < 10:
            raise UploadUnsupported("PDF 无文字层（扫描件）——请粘贴文本或先 OCR")
        return text, source_type

    if ext == ".docx":
        try:
            import docx

            doc = docx.Document(io.BytesIO(raw))
            paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            text = "\n\n".join(paras)
        except Exception as e:
            raise UploadUnsupported(f"docx 解析失败（加密/损坏）：{e}") from e
        if len(text) < 10:
            raise UploadUnsupported("docx 无可提取文本")
        return text, source_type

    raise UploadUnsupported(ext)  # unreachable
CHUNK_OVERLAP_CHARS = 60  # 相邻块重叠，从最近语义边界回退（借鉴 ack-agent chunker）


def _split_sentences(para: str, limit: int) -> list[str]:
    """超长段落按句切（中英文句读），仍超长的句子强制截断。"""
    import re

    parts = re.split(r"(?<=[。；？！.!?;])\s*", para)
    sentences: list[str] = []
    for part in parts:
        while len(part) > limit:
            sentences.append(part[:limit])
            part = part[limit:]
        if part.strip():
            sentences.append(part)
    return sentences


def _overlap_tail(text: str, overlap: int) -> str:
    """取块尾 overlap 字符，并回退到最近语义边界（换行/句读），保证衔接处完整。"""
    tail = text[-overlap:]
    cut = max(tail.rfind("\n"), tail.rfind("。"), tail.rfind("；"), tail.rfind("."))
    return tail[cut + 1 :] if cut > 0 else tail


def split_chunks(content: str, target: int = CHUNK_TARGET_CHARS) -> list[str]:
    """三级分块（借鉴 ack-agent 生产验证）：段落 → 单段超长按句 → 句超长硬截，
    相邻块携带 overlap（语义边界回退），避免边界语义被切断检索不到。"""
    paragraphs = [p.strip() for p in content.split("\n\n") if p.strip()]
    units: list[str] = []
    for para in paragraphs:
        if len(para) <= target:
            units.append(para)
        else:
            units.extend(_split_sentences(para, target))

    chunks: list[str] = []
    buf = ""
    for unit in units:
        if not buf:
            buf = unit
        elif len(buf) + len(unit) + 1 <= target:
            buf = f"{buf}\n{unit}"
        else:
            chunks.append(buf.strip())
            buf = f"{_overlap_tail(buf, CHUNK_OVERLAP_CHARS)}\n{unit}" if CHUNK_OVERLAP_CHARS else unit
    if buf.strip():
        chunks.append(buf.strip())
    return chunks or [content[:target]]


async def create_document(
    db: AsyncSession,
    req: CropDocumentCreate,
    user_id: int | None,
    original_file: tuple[str, str, bytes] | None = None,
) -> CropDocument:
    """吞入+消化：建文档 → 分块 → embedding → 存权威 → 投影到 Redis。

    original_file=(filename, mime, bytes)：上传原件随权威一起落库
    （预览/下载用）；文本粘贴无原件。
    """
    doc = CropDocument(
        title=req.title,
        source_type=req.source_type,
        content=req.content,
        size_bytes=len(req.content.encode("utf-8")),
        status="embedding",
        created_by=user_id,
        **(
            dict(zip(("original_filename", "file_mime", "file_blob"), original_file))
            if original_file
            else {}
        ),
    )
    db.add(doc)
    await db.flush()
    await db.refresh(doc)  # server_default 字段回读（见 01 文档铁坑）

    chunks = split_chunks(req.content)
    try:
        vectors, _mock = await embed_texts(chunks)
        for seq, (text, vec) in enumerate(zip(chunks, vectors)):
            chunk = CropChunk(
                document_id=doc.id,
                seq=seq,
                content=text,
                token_count=len(text),
                embedding=pack_vector(vec),
                embedding_model="mock" if _mock else "api",
            )
            db.add(chunk)
        doc.chunk_count = len(chunks)
        doc.status = "ready"
    except Exception as exc:
        doc.status = "failed"
        doc.error = str(exc)[:500]
        await db.flush()
        raise

    await db.flush()
    # 权威已落库（事务内），投影事务外追加（投影可丢，失败不影响权威）
    rows = (
        await db.execute(
            select(CropChunk.id, CropChunk.embedding).where(
                CropChunk.document_id == doc.id
            )
        )
    ).all()
    for chunk_id, packed in rows:
        await crop_vector_store.add(chunk_id, packed)
    logger.info(
        "crop_ingested", document_id=doc.id, chunks=len(chunks)
    )
    return doc


async def list_documents(
    db: AsyncSession, limit: int = 20, offset: int = 0
) -> tuple[list[CropDocument], int]:
    total = (await db.execute(select(func.count(CropDocument.id)))).scalar_one()
    rows = (
        await db.execute(
            select(CropDocument)
            .order_by(CropDocument.id.desc())
            .limit(limit)
            .offset(offset)
        )
    ).scalars().all()
    return list(rows), total


async def get_document(db: AsyncSession, doc_id: int) -> CropDocument | None:
    return await db.get(CropDocument, doc_id)


async def get_chunks(db: AsyncSession, doc_id: int) -> list[CropChunk]:
    rows = (
        await db.execute(
            select(CropChunk)
            .where(CropChunk.document_id == doc_id)
            .order_by(CropChunk.seq)
        )
    ).scalars().all()
    return list(rows)


async def delete_document(db: AsyncSession, doc_id: int) -> bool:
    """删权威（chunk 级联）+ 删投影。"""
    doc = await get_document(db, doc_id)
    if doc is None:
        return False
    chunk_ids = (
        await db.execute(select(CropChunk.id).where(CropChunk.document_id == doc_id))
    ).scalars().all()
    await db.execute(delete(CropChunk).where(CropChunk.document_id == doc_id))
    await db.delete(doc)
    for cid in chunk_ids:
        await crop_vector_store.drop_chunk(cid)
    return True


async def search(
    db: AsyncSession, query: str, top_k: int = 5
) -> tuple[list[dict], bool]:
    """检索：embed query → Redis KNN → 回 MySQL 取 chunk+文档拼装。"""
    from app.schemas.response import CropSearchHit

    vectors, mock = await embed_texts([query])
    packed = pack_vector(vectors[0])
    hits = await crop_vector_store.search(packed, top_k)
    if not hits:
        return [], mock

    chunk_ids = [cid for cid, _ in hits]
    rows = (
        await db.execute(
            select(CropChunk, CropDocument)
            .join(CropDocument, CropChunk.document_id == CropDocument.id)
            .where(CropChunk.id.in_(chunk_ids))
        )
    ).all()
    by_id = {chunk.id: (chunk, doc) for chunk, doc in rows}
    results: list[CropSearchHit] = []
    for cid, score in hits:  # 保持相似度排序
        pair = by_id.get(cid)
        if pair is None:
            continue  # 投影与权威短暂不一致（重建窗口），跳过
        chunk, doc = pair
        results.append(
            CropSearchHit(
                chunk_id=chunk.id,
                document_id=doc.id,
                document_title=doc.title,
                seq=chunk.seq,
                content=chunk.content,
                score=round(score, 4),
            )
        )
    return [r.model_dump() for r in results], mock


async def rebuild_projection(db: AsyncSession) -> int:
    """从 MySQL 权威全量重建 Redis 投影（运维/灾后备手）。"""
    rows = (
        await db.execute(select(CropChunk.id, CropChunk.embedding))
    ).all()
    return await crop_vector_store.rebuild(rows)


async def projection_health() -> dict:
    """投影健康（管理端观测）。"""
    return {"vector_set": crop_vector_store.VECTOR_KEY, "count": await crop_vector_store.card()}
